import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def create_docx_from_files(directory_path, output_filename):
    """
    从指定目录读取TXT和PNG文件，生成包含内容的DOCX文档
    所有内容使用黑色宋体（中文）和Times New Roman（英文数字），全部居中显示
    """
    doc = Document()

    # 设置全局样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.name = 'Times New Roman'
    style.font.color.rgb = None  # 黑色

    # 初始化主文档标题（稍后会被第一个文本文件的第一行替换）
    main_title = doc.add_paragraph('工程分析报告')
    main_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = main_title.add_run()
    run.bold = True
    run.font.size = Pt(22)

    files = sorted(os.listdir(directory_path))

    # 用于临时存储图片，以便实现并排
    image_queue = []

    # 标志：是否已设置主标题
    main_title_set = False

    def add_images_to_doc():
        """将队列中的图片添加到文档（并排显示）"""
        if not image_queue:
            return

        # 创建一个两列的表格
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置表格宽度为页面宽度
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), '0')  # 0表示100%宽度
        tblPr.append(tblW)

        # 设置单元格边距为0
        for row in table.rows:
            for cell in row.cells:
                # 设置单元格垂直居中
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcPr.append(tcVAlign)

                # 设置单元格边距
                tcMar = OxmlElement('w:tcMar')
                for m in ['top', 'left', 'bottom', 'right']:
                    mar = OxmlElement(f'w:{m}')
                    mar.set(qn('w:w'), '0')
                    mar.set(qn('w:type'), 'dxa')
                    tcMar.append(mar)
                tcPr.append(tcMar)

                # 设置单元格内段落居中
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加图片到表格
        for i, (filename, file_path) in enumerate(image_queue):
            cell_idx = i % 2
            if i % 2 == 0 and i > 0:
                # 添加新行（每两个图片一行）
                table.add_row()

            current_row = table.rows[i // 2]
            cell = current_row.cells[cell_idx]

            # 根据文件名设置图片标题
            display_name = filename  # 默认使用原文件名

            # 根据特定文件名替换为中文标题
            if filename == "all_mises.png":
                display_name = "舱体整体受力"
            elif filename == "frame_deflection.png":
                display_name = "舱体框架挠度"
            elif filename == "frame_mises.png":
                display_name = "舱体框架受力示意图"
            elif filename == "origin_cabin.png":
                display_name = "舱体模型"

            # 添加图片标题
            p = cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(f"{display_name}")
            run.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.name = 'Times New Roman'

            # 添加图片（缩小尺寸）
            p = cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(file_path, width=Inches(2.2))  # 缩小图片宽度

        # 清空图片队列
        image_queue.clear()

        # 添加分隔线
        sep = doc.add_paragraph('-' * 50)
        sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sep.runs[0].font.name = '宋体'
        sep.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        sep.runs[0].font.name = 'Times New Roman'

    def process_table_content(content):
        """处理表格内容，转换为Word表格"""
        # 分割内容为行
        lines = content.split('\n')

        # 存储表格数据
        current_table = []
        table_title = ""
        special_tables = []

        # 用于存储当前处理的表格标题
        current_processing_title = ""

        # 处理特殊表格（钢板用量统计和焊点数量）
        for i, line in enumerate(lines):
            line = line.strip()

            # 检测特殊表格标题
            if any(title in line for title in special_tables):
                # 检查前面是否有空行
                if i > 0 and (not lines[i - 1].strip() or lines[i - 1].strip() == ""):
                    table_title = line
                    # 收集表格数据
                    for j in range(i + 1, len(lines)):
                        data_line = lines[j].strip()
                        if data_line == "":
                            break  # 空行表示表格结束
                        if data_line.startswith('|') and data_line.endswith('|'):
                            # 分割单元格
                            cells = [cell.strip() for cell in data_line[1:-1].split('|')]
                            # 忽略分隔行
                            if not all(re.match(r'^[-+]+$', cell) for cell in cells):
                                current_table.append(cells)

                    # 创建表格
                    if current_table:
                        # 添加表格标题
                        title_para = doc.add_heading(table_title, level=2)
                        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in title_para.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.name = 'Times New Roman'

                        # 创建表格
                        table = doc.add_table(rows=len(current_table), cols=len(current_table[0]))
                        table.style = 'Table Grid'
                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # 设置表格内容居中和字体
                        for row_idx, row_data in enumerate(current_table):
                            row_cells = table.rows[row_idx].cells
                            for col_idx, cell_text in enumerate(row_data):
                                # 设置单元格垂直居中
                                tc = row_cells[col_idx]._tc
                                tcPr = tc.get_or_add_tcPr()
                                tcVAlign = OxmlElement('w:vAlign')
                                tcVAlign.set(qn('w:val'), "center")
                                tcPr.append(tcVAlign)

                                # 设置单元格内容
                                p = row_cells[col_idx].paragraphs[0]
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                run = p.add_run(cell_text)
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                run.font.name = 'Times New Roman'

                        # 重置
                        current_table = []
                        table_title = ""

        # 处理常规表格
        for line in lines:
            line = line.strip()

            # 检测表格标题行（格式：--------------------------标题--------------------------）
            if re.match(r'^-{5,}[^-]+-{5,}$', line):
                # 提取标题文本（去除两边的"-"）
                title_match = re.search(r'^-{5,}([^-]+)-{5,}$', line)
                if title_match:
                    table_title = title_match.group(1).strip()
                    # 存储当前处理的标题
                    current_processing_title = table_title
                continue

            # 检测表格行（以"|"开头和结尾）
            if line.startswith('|') and line.endswith('|'):
                # 分割单元格
                cells = [cell.strip() for cell in line[1:-1].split('|')]

                # 忽略分隔行（只包含"-"的行）
                if all(re.match(r'^[-+]+$', cell) for cell in cells):
                    continue

                # 添加到当前表格
                current_table.append(cells)
            elif current_table:
                # 添加表格标题
                if table_title:
                    title_para = doc.add_heading(table_title, level=2)
                    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in title_para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.name = 'Times New Roman'
                    table_title = ""

                # 创建表格
                table = doc.add_table(rows=len(current_table), cols=len(current_table[0]))
                table.style = 'Table Grid'
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 设置表格内容居中和字体
                for row_idx, row_data in enumerate(current_table):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell_text in enumerate(row_data):
                        # 设置单元格垂直居中
                        tc = row_cells[col_idx]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "center")
                        tcPr.append(tcVAlign)

                        # 设置单元格内容
                        p = row_cells[col_idx].paragraphs[0]
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(cell_text)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.name = 'Times New Roman'

                # 重置当前表格
                current_table = []

        # 处理最后一个表格
        if current_table:
            # 添加表格标题
            if table_title:
                title_para = doc.add_heading(table_title, level=2)
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in title_para.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.name = 'Times New Roman'

            # 创建表格
            if current_table:
                table = doc.add_table(rows=len(current_table), cols=len(current_table[0]))
                table.style = 'Table Grid'
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 设置表格内容居中和字体
                for row_idx, row_data in enumerate(current_table):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell_text in enumerate(row_data):
                        # 设置单元格垂直居中
                        tc = row_cells[col_idx]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = OxmlElement('w:vAlign')
                        tcVAlign.set(qn('w:val'), "center")
                        tcPr.append(tcVAlign)

                        # 设置单元格内容
                        p = row_cells[col_idx].paragraphs[0]
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = p.add_run(cell_text)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.name = 'Times New Roman'

    for filename in files:
        file_path = os.path.join(directory_path, filename)

        if filename.lower().endswith('.txt'):
            # 处理文本文件前，先处理队列中的图片
            if image_queue:
                add_images_to_doc()

            try:
                # 读取文本内容
                with open(file_path, 'r', encoding='utf-8') as txt_file:
                    content_lines = txt_file.readlines()

                # 使用第一行作为文档标题
                if content_lines:
                    first_line = content_lines[0].strip()

                    # 如果是第一个文本文件，设置为主标题
                    if not main_title_set:
                        # 更新主标题
                        for paragraph in doc.paragraphs:
                            if '工程分析报告' in paragraph.text:
                                # 清除原有内容
                                for run in paragraph.runs:
                                    run.text = ""
                                # 添加新内容
                                run = paragraph.add_run(first_line)
                                run.bold = True
                                run.font.size = Pt(22)
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                run.font.name = 'Times New Roman'
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                break
                        main_title_set = True
                    else:
                        # 添加文件标题
                        heading = doc.add_heading(first_line, level=1)
                        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in heading.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.name = 'Times New Roman'

                    # 处理剩余内容
                    remaining_content = ''.join(content_lines[1:])
                    # 处理表格内容
                    process_table_content(remaining_content)

                # 添加分隔线
                sep = doc.add_paragraph('-' * 50)
                sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                sep.runs[0].font.name = '宋体'
                sep.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                sep.runs[0].font.name = 'Times New Roman'

            except Exception as e:
                print(f"处理文本文件 {filename} 时出错: {str(e)}")
        elif filename.lower().endswith('.png'):
            # 将图片添加到队列
            image_queue.append((filename, file_path))

    # 处理剩余的图片
    if image_queue:
        add_images_to_doc()

    # 确保所有内容居中
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 保存文档
    save_path = os.path.join(directory_path, output_filename)
    doc.save(save_path)  # type: ignore[reportGeneralTypeIssues]