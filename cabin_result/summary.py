import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
import sys
import copy
from docx.enum.table import WD_TABLE_ALIGNMENT


def copy_docx_images(src_doc, dest_doc):
    """复制源文档中的所有图片到目标文档"""
    # 获取源文档的图片关系
    src_rel = src_doc.part.rels
    img_rels = [r for r in src_rel.values() if "image" in r.target_ref]

    # 复制图片到目标文档
    for rel in img_rels:
        img_data = rel.target_part.blob
        dest_doc.part.related_parts[rel.rId] = rel.target_part
        dest_doc.part.rels.add_relationship(
            rel.reltype,
            rel.target_part,
            rel.rId,
            is_external=rel.is_external
        )

def merge_documents_with_images(docs):
    """合并多个文档并保留图片"""
    if not docs:
        return None

    # 创建新文档作为基础
    base_doc = docs[0]

    # 复制图片
    for doc in docs:
        if doc != base_doc:
            copy_docx_images(doc, base_doc)

    # 合并文档内容
    for doc in docs[1:]:
        for element in doc.element.body:
            new_element = copy.deepcopy(element)
            base_doc.element.body.append(new_element)

    return base_doc

def remove_leading_breaks(doc):
    """移除文档开头不必要的分页符和空段落"""
    # 查找并移除开头的分页符
    for para in doc.paragraphs:
        # 检查段落是否包含分页符
        runs = para.runs
        if runs and runs[0].text == '' and runs[0]._element.xml.find('w:br') != -1:
            # 移除这个空段落
            p = para._element
            p.getparent().remove(p)
        elif para.text.strip() != '':
            # 遇到有内容的段落，停止处理
            break

def extract_folder_info(folder_path):
    """从文件夹路径中提取关键信息"""
    # 获取文件夹名称
    folder_name = os.path.basename(folder_path)
    # print(f"  文件夹名称: {folder_name}")

    # 提取数字信息
    match = re.search(r'iter_(\d+\.\d+)_(\d+)_', folder_name)
    if match:
        size = match.group(1)  # 500.0
        iteration = int(match.group(2))  # 转换为整数
        # print(f"  提取尺寸: {size}mm, 迭代次数: {iteration}")
        return size, iteration
    else:
        # print("  无法从文件夹名称提取信息")
        return None, None

def modify_title(doc, size, iteration):
    """修改文档标题为宋体2号字体"""
    if size and iteration is not None:
        # 创建新标题文本
        new_title = f"{size}mm，第{iteration}次迭代结果"

        # 修改第一段内容
        if len(doc.paragraphs) > 0:
            first_para = doc.paragraphs[0]

            # 清除原有内容
            first_para.clear()

            # 添加新内容并设置格式
            runner = first_para.add_run(new_title)
            runner.font.name = '宋体'  # 设置字体为宋体
            runner._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 确保中文字体生效
            runner.font.size = Pt(22)  # 2号字对应22磅
            runner.bold = True  # 加粗

            # 设置段落居中
            first_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # print(f"  已修改标题: {new_title} (宋体2号居中)")
        else:
            pass
            # print("  文档没有段落，无法修改标题")
        return new_title
    return None

def extract_total_weight(doc):
    """从文档中提取钢材总重量"""
    total = 0.0

    # 直接定位特定表格
    tables = doc.tables

    # 表格10: 型材重量数据 (索引9)
    if len(tables) > 9:
        table10 = tables[9]
        # print(f"  处理型材表格 (表格10): {len(table10.rows)}行 x {len(table10.columns)}列")

        # 提取最后一列数据
        for row_idx in range(0, len(table10.rows)):
            row = table10.rows[row_idx]
            if len(row.cells) >= 4:  # 确保有4列
                cell_text = row.cells[3].text.strip()
                if cell_text and cell_text.replace('.', '', 1).isdigit():  # 检查是否为数字
                    try:
                        weight = float(cell_text)
                        total += weight
                        # print(f"    添加型材重量: {weight:.2f}kg (行{row_idx + 1})")
                    except ValueError:
                        pass
                        # print(f"    解析错误: '{cell_text}'")

    # 表格12: 钢板重量数据 (索引11)
    if len(tables) > 11:
        table12 = tables[11]
        # print(f"  处理钢板表格 (表格12): {len(table12.rows)}行 x {len(table12.columns)}列")

        # 提取最后一列数据
        for row_idx in range(0, len(table12.rows)):
            row = table12.rows[row_idx]
            if len(row.cells) >= 4:  # 确保有4列
                cell_text = row.cells[3].text.strip()
                if cell_text and cell_text.replace('.', '', 1).isdigit():  # 检查是否为数字
                    try:
                        weight = float(cell_text)
                        total += weight
                        # print(f"    添加钢板重量: {weight:.2f}kg (行{row_idx + 1})")
                    except ValueError:
                        pass
                        # print(f"    解析错误: '{cell_text}'")

    return total

def add_total_weight_info(doc: Document, total_weight: float) -> float:
    """在文档末尾添加总重信息并分页，设置中文为黑体，西文为Times New Roman"""
    if total_weight > 0:
        # 添加总重信息段落
        total_para = doc.add_paragraph()
        total_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 创建加粗文本并设置字体
        runner = total_para.add_run(f"钢材用量总重为{total_weight:.2f}kg")
        runner.font.size = Pt(14)
        runner.bold = True
        # 设置西文字体
        runner.font.name = 'Times New Roman'
        # 设置中文字体
        r_fonts = runner._element.rPr.rFonts
        r_fonts.set(qn('w:eastAsia'), 'SimHei')

        # 添加分页符 - 只在最后一个方案不添加
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # print(f"  已添加总重信息和分页符")
        return total_weight
    return 0.0

def add_comparison_table(final_doc, summary_data):
    """在文档开头添加方案对比表格"""
    if not summary_data:
        return

    # 添加汇总报告标题
    title_para = final_doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("汇总报告")
    title_run.font.name = '宋体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_run.font.size = Pt(22)  # 二号字
    title_run.bold = True

    # 添加空行
    final_doc.add_paragraph()

    # 添加方案对比标题
    subtitle_para = final_doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle_para.add_run("各方案钢材用量对比")
    subtitle_run.font.name = '宋体'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    subtitle_run.font.size = Pt(16)
    subtitle_run.bold = True

    # 添加空行
    final_doc.add_paragraph()

    # 创建表格 (行数 = 数据行数 + 表头)
    table = final_doc.add_table(rows=len(summary_data) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格样式
    table.style = 'Table Grid'

    # 添加表头
    header_cells = table.rows[0].cells
    header_cells[0].text = "方案"
    header_cells[1].text = "钢材总用量 (kg)"

    # 设置表头格式
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 找到最小用钢量的方案
    min_weight = min(data[1] for data in summary_data)

    # 填充表格数据
    for i, (title, weight) in enumerate(summary_data):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = title
        row_cells[1].text = f"{weight:.2f}"

        # 设置单元格字体为宋体
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 如果是用钢量最小的方案，标记为红色
        if weight == min_weight:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # 红色

    # 在表格后添加空行
    final_doc.add_paragraph()

    # print("\n已添加方案对比表格")

def process_single_docx(input_folder):
    """处理单个文件夹中的文档"""
    # 提取文件夹信息
    size, iteration = extract_folder_info(input_folder)

    # 如果无法提取文件夹信息，跳过处理
    if size is None or iteration is None:
        return None, None, None

    # 查找文件夹中的docx文件
    docx_files = [f for f in os.listdir(input_folder)
                  if f.endswith('.docx') and not f.startswith('~$')]

    if not docx_files:
        # print(f"  在文件夹 {input_folder} 中未找到docx文件")
        return None, None, None

    # 处理每个文档
    for filename in docx_files:
        input_path = os.path.join(input_folder, filename)

        # print(f"\n处理文件: {filename}")

        try:
            # 打开文档
            doc = Document(input_path)

            # 修改标题并获取标题文本
            title = modify_title(doc, size, iteration)

            # 提取总重量
            total_weight = extract_total_weight(doc)

            if total_weight > 0:
                # print(f"  钢材总重量: {total_weight:.2f}kg")

                # 添加总重信息
                add_total_weight_info(doc, total_weight)

                # 移除文档开头不必要的分页符
                remove_leading_breaks(doc)

                return title, total_weight, doc
            else:
                # print("  未提取到有效重量数据")
                return None, None, None

        except Exception as e:
            # print(f"  处理失败: {str(e)}")
            # traceback.# print_exc()
            return None, None, None
    return None, None, None

def filter_folders_by_max_iteration(folders, main_folder):
    """
    筛选文件夹：对于每个尺寸，只保留迭代次数最大的文件夹
    返回筛选后的文件夹列表
    """
    # 使用字典按尺寸分组
    size_groups = {}

    # 收集所有文件夹信息
    for folder in folders:
        folder_path = os.path.join(main_folder, folder)
        size, iteration = extract_folder_info(folder_path)

        # 确保提取成功
        if size and iteration is not None:
            # 添加到对应尺寸的分组
            if size not in size_groups:
                size_groups[size] = []
            size_groups[size].append((iteration, folder))

    # 筛选每个尺寸的最大迭代文件夹
    selected_folders = []

    for size, folders_list in size_groups.items():
        # 按迭代次数降序排序
        folders_list.sort(key=lambda x: x[0], reverse=True)

        # 取迭代次数最大的文件夹
        max_iteration, folder_name = folders_list[0]
        selected_folders.append(folder_name)

        # print(f"尺寸 {size}mm: 选择迭代次数最大的文件夹 {folder_name} (迭代次数: {max_iteration})")

    return selected_folders

def process_all_folders(main_folder):
    """处理主文件夹中的所有迭代文件夹"""
    # 创建汇总文档
    final_doc = Document()

    # 设置默认字体为宋体
    for style in final_doc.styles:
        if style.type == 1:  # 段落样式
            style.font.name = '宋体'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 查找所有迭代文件夹
    iter_folders = [f for f in os.listdir(main_folder)
                    if os.path.isdir(os.path.join(main_folder, f)) and f.startswith('iter_')]

    if not iter_folders:
        # print(f"在文件夹 {main_folder} 中未找到迭代文件夹")
        return

    # print(f"找到 {len(iter_folders)} 个迭代文件夹:")
    for folder in iter_folders:
        pass
        # print(f"  - {folder}")

    # 筛选文件夹：每个尺寸只保留迭代次数最大的文件夹
    selected_folders = filter_folders_by_max_iteration(iter_folders, main_folder)

    if not selected_folders:
        # print("没有符合条件的文件夹需要处理")
        return

    # print(f"\n筛选后需要处理 {len(selected_folders)} 个文件夹:")
    for folder in selected_folders:
        pass
        # print(f"  - {folder}")

    # 按文件夹名称排序
    selected_folders.sort()

    # 存储所有方案的数据（标题，总重量）
    summary_data = []

    # 存储所有处理后的文档对象
    processed_docs = []

    # 处理每个筛选后的迭代文件夹
    for i, folder in enumerate(selected_folders):
        folder_path = os.path.join(main_folder, folder)
        # print(f"\n{'=' * 50}")
        # print(f"处理迭代文件夹: {folder} ({i + 1}/{len(selected_folders)})")

        # 处理文件夹中的文档
        title, total_weight, processed_doc = process_single_docx(folder_path)

        if not processed_doc:
            pass
            # print(f"  文件夹 {folder} 处理失败，跳过")
        else:
            # 保存方案数据
            summary_data.append((title, total_weight))
            processed_docs.append(processed_doc)

    # 合并所有文档并保留图片
    if processed_docs:
        # 创建临时文档用于添加表格
        temp_doc = Document()

        # 添加对比表格
        add_comparison_table(temp_doc, summary_data)

        # 将临时文档作为第一个文档
        all_docs = [temp_doc] + processed_docs

        # 合并所有文档（包括对比表格和各方案文档）
        final_doc = merge_documents_with_images(all_docs)

        # 移除汇总文档开头的空白页
        remove_leading_breaks(final_doc)

        # 移除最后一个文档末尾的分页符
        if final_doc.paragraphs:
            last_para = final_doc.paragraphs[-1]
            if last_para.runs and last_para.runs[-1].text == '':
                last_para._element.getparent().remove(last_para._element)

    # 保存汇总文档
    output_path = os.path.join(main_folder, "方案比选报告.docx")
    final_doc.save(output_path)

    return output_path
    # print(f"\n{'=' * 50}")
    # print(f"处理完成! 汇总文档已保存至: {output_path}")
    # print(f"共处理了 {len(selected_folders)} 个迭代文件夹 (每个尺寸的最新迭代)")

if __name__ == "__main__":
    # 检查是否提供了主文件夹路径
    if len(sys.argv) > 1:
        main_folder = sys.argv[1]
    else:
        main_folder = input("请输入包含迭代文件夹的主文件夹路径: ")

    # 确保路径存在
    if not os.path.exists(main_folder):
        # print(f"错误: 路径 '{main_folder}' 不存在")
        exit(1)

    # 处理所有文件夹
    process_all_folders(main_folder)